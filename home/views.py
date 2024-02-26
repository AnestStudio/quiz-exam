import ast
import copy
import datetime
import json
import os
import random
import string
from datetime import datetime

from django.conf import settings
from django.contrib.auth import logout
from django.contrib.auth.decorators import login_required
from django.contrib.auth.models import User
from django.contrib.auth.views import LoginView, PasswordChangeView, PasswordResetView, PasswordResetConfirmView
from django.db.models import OuterRef, Subquery
from django.http import JsonResponse
from django.shortcuts import render, redirect, get_object_or_404
from django.urls import reverse
from docx import Document as DocxDocument

from home.forms import RegistrationForm, UserLoginForm, UserPasswordResetForm, UserPasswordChangeForm, \
    UserSetPasswordForm
from .forms import DocumentForm
from .models import Question, Subject, Exam, ExamQuestion, SubjectExam, ExamSchedule


# Pages
def index(request):
    return render(request, 'pages/index.html')


class UserLoginView(LoginView):
    template_name = 'accounts/sign-in.html'
    form_class = UserLoginForm


def logout_view(request):
    logout(request)
    return redirect('/accounts/login')


def register(request):
    if request.method == 'POST':
        form = RegistrationForm(request.POST)
        if form.is_valid():
            form.save()
            print("Account created successfully!")
            return redirect('/accounts/login')
        else:
            print("Registration failed!")
    else:
        form = RegistrationForm()
    return render(request, 'accounts/sign-up.html', {'form': form})


class UserPasswordResetView(PasswordResetView):
    template_name = 'accounts/password_reset.html'
    form_class = UserPasswordResetForm


class UserPasswordResetConfirmView(PasswordResetConfirmView):
    template_name = 'accounts/password_reset_confirm.html'
    form_class = UserSetPasswordForm


class UserPasswordChangeView(PasswordChangeView):
    template_name = 'accounts/password_change.html'
    form_class = UserPasswordChangeForm


# IMPORT QUESTION ------------------------------------------------------------------------------------------------------

@login_required
def upload_document(request):
    if request.method == 'POST':
        form = DocumentForm(request.POST, request.FILES)
        if form.is_valid():
            docx_file = request.FILES['file']
            data = __read_data_documents(docx_file)

            # Clone dict
            data_insert = copy.deepcopy(data)
            for e in data_insert['questions']:
                e['answer_list'] = str(e['answer_list']).replace("\'", "\"")
                e['answer'] = str(e['answer']).replace("\'", "\"")

            __insert_question_list_to_db(data_insert['questions'])

            return render(
                request,
                'pages/question/import-question-success.html',
                {
                    'message': 'Import question successfully!',
                    'data': {
                        'subject': data['subject'],
                        'number_of_quiz': data['number_of_quiz'],
                        'lecturer': data['lecturer'],
                        'date': data['date'],
                        'questions': data['questions']
                    },
                    'page': 11
                }
            )
    else:
        form = DocumentForm()
    return render(request, 'pages/question/import-question.html', {'form': form, 'page': 11})


def __read_data_documents(file):
    document = DocxDocument(file)
    data = dict()

    for paragraph in document.paragraphs:
        result = paragraph.text.strip().split(':')
        if len(result) == 2:
            data[result[0].strip().replace(' ', '_').lower()] = result[1].strip()

    list_question = list()
    for table in document.tables:
        question = dict()

        # Get question
        question['question'] = table.rows[0].cells[1].text.strip()

        # Get list answer of question
        answer_list = dict()
        for i in range(1, len(table.rows)):
            if table.rows[i].cells[0].text.strip() == 'ANSWER:':
                break

            answer_list[table.rows[i].cells[0].text.strip()[:1]] = table.rows[i].cells[1].text.strip()
            question['answer_list'] = answer_list

        for i in range(len(answer_list) + 1, len(table.rows)):
            temp = table.rows[i].cells[0].text.strip()
            key = temp[0: len(temp) - 1].replace(" ", "_").lower()
            value = table.rows[i].cells[1].text.strip().lower()

            if temp == 'MIX CHOICES:':
                question[key] = True if value == 'yes' else False
            elif temp == 'ANSWER:':
                if len(value) == 1:
                    question[key] = list(value)
                else:
                    question[key] = value.replace(' ', '').split(',')
            else:
                question[key] = value

        # TODO: hardcode subject_id
        question['subject'] = Subject.objects.get(subject_id=2)
        question['image'] = 'no_image'

        # Add question to list
        list_question.append(question)

    data['questions'] = list_question
    return data


def __insert_question_list_to_db(data_list):
    for data_dict in data_list:
        data_model_instance = Question(**data_dict)
        data_model_instance.save()


# MANAGE QUESTION ------------------------------------------------------------------------------------------------------

@login_required
def view_list_question(request):
    subjects = __get_subject_list_by_user(request.user.id)
    subject_id = 0
    questions = None

    if request.method == 'POST':
        subject = request.POST.get('subject')
        if subject:
            if subject == '0':
                questions = ''
            else:
                questions = Question.objects.filter(subject=subject)
                questions = sorted(questions, key=lambda o: o.unit)
                subject_id = int(subject)
    else:
        questions = ''

    data = {
        'subject_id': subject_id,
        'subjects': subjects,
        'questions': questions,
        'page': 21
    }
    return render(request, 'pages/question/question.html', data)


@login_required
def view_question(request, question_id):
    subjects = __get_subject_list_by_user(request.user.id)
    question = Question.objects.get(question_id=question_id)

    # Convert the string list answer to a dictionary
    question.answer_list = eval(question.answer_list)

    data = {
        'subject_id': int(question.subject.subject_id),
        'subjects': subjects,
        'question': question,
        'page': 21
    }
    return render(request, 'pages/question/view-question.html', data)


@login_required
def add_new_question(request):
    subjects = __get_subject_list_by_user(request.user.id)

    if request.method == 'POST':
        subject_id = request.POST.get('subject')
        question_content = request.POST.get('question_content')
        question_unit = request.POST.get('question_unit')
        mix_choices = request.POST.get('mix_choices')
        answer_list = request.POST.getlist('answers')
        answer = request.POST.getlist('correct_answer')
        question_image = 'no_image'

        if 'question_image' in request.FILES:
            image_file = request.FILES['question_image']
            question_image = image_file.name

            upload_dir = os.path.join(settings.MEDIA_ROOT, 'images')
            if not os.path.exists(upload_dir):
                os.makedirs(upload_dir)

            with open(os.path.join(upload_dir, question_image), 'wb+') as destination:
                for chunk in image_file.chunks():
                    destination.write(chunk)

        question = Question(
            subject_id=subject_id,
            question=question_content,
            answer_list=str(__convert_answer_list_to_dict(answer_list)).replace("\'", "\""),
            answer=str(answer).replace("\'", "\""),
            mark=0.5,
            unit=question_unit,
            mix_choices=True if mix_choices == '1' else False,
            image=question_image
        )
        question.save()

        # Convert the string list answer to a dictionary
        question.answer_list = ast.literal_eval(question.answer_list)
        data = {
            'question': question,
            'add_new_question_success': True,
            'page': 22
        }
        return render(request, 'pages/question/view-question.html', data)
    else:
        data = {
            'subjects': subjects,
            'page': 22
        }
        return render(request, 'pages/question/add-new-question.html', data)


@login_required
def delete_question(request, question_id):
    subjects = __get_subject_list_by_user(request.user.id)
    question = Question.objects.get(question_id=question_id)
    questions = Question.objects.filter(subject=question.subject.subject_id)

    Question.objects.filter(question_id=question_id).delete()

    data = {
        'subject_id': int(question.subject.subject_id),
        'subjects': subjects,
        'questions': questions,
        'delete_question_success': True,
        'page': 21
    }
    return render(request, 'pages/question/question.html', data)


@login_required
def check_question_in_exam(request, question_id):
    try:
        exam_questions = ExamQuestion.objects.filter(question_id=question_id).all()
        exams = []
        for exam_question in exam_questions:
            exams.append(exam_question.exam.exam_code)
        response = {}
        if len(exam_questions) > 0:
            response['length'] = len(exam_questions)
            response['exams'] = exams
        else:
            response['length'] = 0
        return JsonResponse(response, safe=False)
    except ExamQuestion.DoesNotExist:
        return JsonResponse({'error': 'Object not found'}, status=404)


@login_required
def view_edit_question(request, question_id):
    subjects = __get_subject_list_by_user(request.user.id)
    question = Question.objects.get(question_id=question_id)

    # Convert the string list answer to a dictionary
    question.answer_list = ast.literal_eval(question.answer_list)

    data = {
        'subject_id': int(question.subject.subject_id),
        'subjects': subjects,
        'question': question,
        'page': 21
    }
    return render(request, 'pages/question/edit-question.html', data)


@login_required
def edit_question(request):
    if request.method == 'POST':
        subject_id = request.POST.get('subject')
        question_id = request.POST.get('question_id')
        question_content = request.POST.get('question_content')
        question_unit = request.POST.get('question_unit')
        mix_choices = request.POST.get('mix_choices')
        answer_list = request.POST.getlist('answers')
        answer = request.POST.getlist('correct_answer')
        question_image = 'no_image'

        if 'question_image' in request.FILES:
            image_file = request.FILES['question_image']
            question_image = image_file.name

            upload_dir = os.path.join(settings.MEDIA_ROOT, 'images')
            if not os.path.exists(upload_dir):
                os.makedirs(upload_dir)

            with open(os.path.join(upload_dir, question_image), 'wb+') as destination:
                for chunk in image_file.chunks():
                    destination.write(chunk)

        # Fetch the question want to update
        question = get_object_or_404(Question, question_id=question_id)
        # Modify the fields of the question
        question.subject_id = subject_id
        question.question = question_content
        question.answer_list = str(__convert_answer_list_to_dict(answer_list)).replace("\'", "\"")
        question.answer = str(answer).replace("\'", "\"")
        print(question.answer_list)
        print(question.answer)
        question.unit = question_unit
        question.mix_choices = True if mix_choices == '1' else False
        question.image = question_image
        question.save()

        # Sau khi update correct answer ở bảng [question] thì cần update lại correct answer bảng [exam_question]
        exam_questions = ExamQuestion.objects.filter(question_id=question_id)
        for exam_question in exam_questions:
            dict_answer_random = __convert_answer_list_to_dict(__random_dict_value_position(eval(question.answer_list)))
            list_correct_answer_new = __get_correct_list_answer_new(
                eval(question.answer_list),
                dict_answer_random,
                ast.literal_eval(question.answer)
            )
            exam_question.answer_list = str(dict_answer_random).replace("\'", "\"")
            exam_question.answer = str(list_correct_answer_new).replace("\'", "\"")
            exam_question.save()

        # Convert the string list answer to a dictionary
        question.answer_list = ast.literal_eval(question.answer_list)

        data = {
            'question': question,
            'update_question_success': True,
            'page': 21
        }
        return render(request, 'pages/question/view-question.html', data)
    else:
        return render(request, 'pages/404.html')


# MANAGE EXAM ----------------------------------------------------------------------------------------------------------

@login_required
def view_list_exam(request):
    subjects = __get_subject_list_by_user(request.user.id)
    subject_id = 0
    exams = None

    if request.method == 'POST':
        subject_id = request.POST.get('subject')
        if subject_id:
            if subject_id == '0':
                exams = ''
            else:
                exams = __get_exam_list_by_subject(subject_id)
    else:
        exams = ''

    data = {
        'subject_id': int(subject_id),
        'subjects': subjects,
        'exams': exams,
        'page': 31
    }
    return render(request, 'pages/exam/exam.html', data)


@login_required
def view_exam(request, exam_id):
    data = __go_to_view_exam(exam_id)
    return render(request, 'pages/exam/view-exam.html', data)


@login_required
def view_question_in_exam(request, exam_id, question_id):
    try:
        exam_question = ExamQuestion.objects.get(exam_id=exam_id, question_id=question_id)

        # Convert the object to a dictionary
        exam_question_dict = {
            'question_id': exam_question.question.question_id,
            'question': exam_question.question.question,
            'answer_list': json.loads(exam_question.answer_list.replace("\'", "\"")),
            'answer': exam_question.answer.replace("\"", "\'"),
            'mark': str(exam_question.question.mark),
            'unit': exam_question.question.unit,
            'mix_choices': exam_question.question.mix_choices,
            'image': exam_question.question.image.name,
            'subject': {
                'subject_id': exam_question.question.subject.subject_id,
                'subject_code': exam_question.question.subject.subject_code,
                'subject_name': exam_question.question.subject.subject_name
            }
        }
        # Return JSON response
        return JsonResponse(exam_question_dict, safe=False)

    except ExamQuestion.DoesNotExist:
        return JsonResponse({'error': 'Object not found'}, status=404)


@login_required
def delete_question_in_exam(request, exam_id, question_id):
    ExamQuestion.objects.filter(exam_id=exam_id, question_id=question_id).delete()
    data = __go_to_view_exam(exam_id)
    data['delete_question_in_exam_success'] = True
    return render(request, 'pages/exam/view-exam.html', data)


def __go_to_view_exam(exam_id):
    subject_exam = SubjectExam.objects.get(exam_id=exam_id)
    list_question_temp = sorted(__get_question_list_by_exam(exam_id), key=lambda o: o.unit)

    questions = []
    for question in list_question_temp:
        question.answer_list = eval(question.answer_list)
        questions.append(question)

    return {
        'subject_exam': subject_exam,
        'questions': questions,
        'page': 31
    }


@login_required
def view_generate_exam(request):
    subjects = __get_subject_list_by_user(request.user.id)
    data = {}

    if request.method == 'POST':
        subject_id = request.POST.get('subject')
        if subject_id:
            data = __go_to_view_generate_exam(subject_id)
    else:
        data['chapters'] = ''

    data['subjects'] = subjects
    data['page'] = 32
    return render(request, 'pages/exam/generate-exam.html', data)


@login_required
def generate_exam(request):
    if request.method == 'POST':
        subject_id = request.POST.get('subject_id')
        exam_code = request.POST.get('exam_code')
        exam_duration = request.POST.get('exam_duration')
        number_question = request.POST.get('number_question')
        list_question_id = request.POST.getlist('question')
        list_chapter_name = request.POST.getlist('chapter')

        list_question_reselect = Question.objects.filter(question_id__in=list_question_id)
        list_question_by_chapter = Question.objects.filter(subject=subject_id, unit__in=list_chapter_name)
        list_question_generate = [q for q in list_question_by_chapter if q not in list_question_reselect]

        number_of_question_random = int(number_question) - len(list_question_reselect)
        if number_of_question_random > len(list_question_generate):
            data = __go_to_view_generate_exam(subject_id)
            data['subjects'] = __get_subject_list_by_user(request.user.id)
            data['error_size'] = True
            data['page'] = 32
            return render(request, 'pages/exam/generate-exam.html', data)
        else:
            random_list_question = random.sample(list_question_generate, number_of_question_random)
            exam_insert = Exam.objects.create(
                user_id=request.user.id,
                exam_code=exam_code,
                question_count=number_question,
                exam_duration=exam_duration,
                exam_score=10
            )
            SubjectExam.objects.create(
                exam_id=exam_insert.exam_id,
                subject_id=subject_id
            )
            for q in (list(list_question_reselect) + random_list_question):
                if q.mix_choices:
                    dict_answer_random = __convert_answer_list_to_dict(__random_dict_value_position(eval(q.answer_list)))
                    list_correct_answer_new = __get_correct_list_answer_new(
                        eval(q.answer_list),
                        dict_answer_random,
                        ast.literal_eval(q.answer)
                    )
                else:
                    dict_answer_random = q.answer_list
                    list_correct_answer_new = q.answer

                ExamQuestion.objects.create(
                    exam_id=exam_insert.exam_id,
                    question_id=q.question_id,
                    answer_list=str(dict_answer_random).replace("\'", "\""),
                    answer=str(list_correct_answer_new).replace("\'", "\"")
                )
        return redirect(reverse('view_exam', kwargs={'exam_id': exam_insert.exam_id}))
    else:
        return render(request, 'pages/404.html')


@login_required
def edit_exam(request, exam_id):
    return render(request, 'pages/exam/exam.html')


@login_required
def check_exam_schedule(request, exam_id):
    try:
        exam_schedule = ExamSchedule.objects.filter(exam_id=exam_id).all()
        response = {}
        if len(exam_schedule) > 0:
            response['length'] = len(exam_schedule)
        else:
            response['length'] = 0
        # Return JSON response
        return JsonResponse(response, safe=False)
    except ExamQuestion.DoesNotExist:
        return JsonResponse({'error': 'Object not found'}, status=404)


@login_required
def delete_exam(request, subject_id, exam_id):
    Exam.objects.filter(exam_id=exam_id).delete()
    subjects = __get_subject_list_by_user(request.user.id)
    exams = __get_exam_list_by_subject(subject_id)

    data = {
        'subject_id': int(subject_id),
        'subjects': subjects,
        'exams': exams,
        'delete_exam_success': True,
        'page': 31
    }
    return render(request, 'pages/exam/exam.html', data)


def __go_to_view_generate_exam(subject_id):
    if subject_id == '0':
        chapters = ''
        questions = []
    else:
        # Get all questions => list chapter
        list_question_all = Question.objects.filter(subject=subject_id)
        chapters = list()
        for q in list_question_all:
            chapters.append(q.unit)
        chapters = sorted(list(set(chapters)))

        questions_generated = ExamQuestion.objects.filter(
            id=Subquery(
                ExamQuestion.objects.filter(question_id=OuterRef('question_id')).order_by('id').values('id')[:1]
            )
        )
        questions_generated_chapter = []
        for q in questions_generated:
            if q.question.subject_id == int(subject_id):
                questions_generated_chapter.append(q.question)

        questions = sorted(questions_generated_chapter, key=lambda o: o.unit)
        subject_id = int(subject_id)

    return {
        'subject_id': subject_id,
        'questions': questions,
        'chapters': chapters,
    }


# EXAM SCHEDULE --------------------------------------------------------------------------------------------------------

@login_required
def schedule_exam(request):
    subjects = __get_subject_list_by_user(request.user.id)
    subject_id = 0
    exams = []
    add_success = False

    if request.method == 'POST':
        if 'save' in request.POST:
            exam_id = request.POST.get('exam_id')
            user_id = request.user.id

            start_date_get = request.POST.get('start_date')
            end_date_get = request.POST.get('end_date')

            exam_schedule = ExamSchedule.objects.create(exam_id=exam_id, user_id=user_id,
                                                        start_time=datetime.strptime(start_date_get, '%Y-%m-%dT%H:%M'),
                                                        end_time=datetime.strptime(end_date_get, '%Y-%m-%dT%H:%M'))
            exam_schedule.save()
            exam_schedules = ExamSchedule.objects.filter(exam_id=exam_schedule.exam_id)
            subject_exam = SubjectExam.objects.get(exam_id=exam_schedules[0].exam.exam_id)
            data = {
                'exam_schedules': exam_schedules,
                'subject_exam': subject_exam,
                'add_success': True,
                'page': 42
            }
            return render(request, 'pages/schedule/view_detail_schedule.html', data)
        elif 'subject' in request.POST:
            subject_id = request.POST.get('subject')
            if subject_id:
                exams = __get_exam_list_by_subject(subject_id)
        else:
            return render(request, 'pages/404.html')

    data = {
        'subject_id': int(subject_id),
        'subjects': subjects,
        'exams': exams,
        'add_success': add_success,
        'page': 41
    }
    return render(request, 'pages/schedule/schedule-exam.html', data)


@login_required
def view_schedule_exam(request):
    subject_id = 0
    exams = None
    subjects = __get_subject_list_by_user(request.user.id)
    if request.method == 'POST':
        if 'subject' in request.POST:
            subject_id = request.POST.get('subject')
            if subject_id:
                if subject_id == '0':
                    exams = ''
                else:
                    exams = __get_exam_have_schedule_list_by_subject(subject_id)
    else:
        exams = ''

    data = {
        'subject_id': int(subject_id),
        'subjects': subjects,
        'exams': exams,
        'page': 42
    }
    return render(request, 'pages/schedule/view-schedule-exam.html', data)


@login_required
def view_detail_schedule_exam(request, exam_id):
    exam_schedules = ExamSchedule.objects.filter(exam_id=exam_id)
    subject_exam = SubjectExam.objects.get(exam_id=exam_schedules[0].exam.exam_id)

    data = {
        'exam_schedules': exam_schedules,
        'subject_exam': subject_exam,
        'page': 42
    }
    return render(request, 'pages/schedule/view_detail_schedule.html', data)


@login_required
def view_edit_schedule_exam(request, schedule_exam_id):
    exam_schedule = ExamSchedule.objects.get(schedule_id=schedule_exam_id)

    subject_exam = SubjectExam.objects.get(exam_id=exam_schedule.exam.exam_id)
    subject_name = subject_exam.subject.subject_name

    data = {
        'exam_schedule': exam_schedule,
        'subject_name': subject_name,
        'page': 42
    }
    return render(request, 'pages/schedule/edit-schedule-exam.html', data)


@login_required
def edit_schedule_exam(request):
    if request.method == 'POST':
        schedule_exam_id = request.POST.get('schedule_exam_id')
        start_date = request.POST.get('start_date')
        end_date = request.POST.get('end_date')

        # Get one exam schedule
        exam_schedule = ExamSchedule.objects.get(schedule_id=int(schedule_exam_id))

        exam_schedule.start_time = datetime.strptime(start_date, '%Y-%m-%dT%H:%M')
        exam_schedule.end_time = datetime.strptime(end_date, '%Y-%m-%dT%H:%M')
        exam_schedule.save()

        exam_schedules = ExamSchedule.objects.filter(exam_id=exam_schedule.exam_id)
        subject_exam = SubjectExam.objects.get(exam_id=exam_schedules[0].exam.exam_id)

        data = {
            'exam_schedules': exam_schedules,
            'subject_exam': subject_exam,
            'update_schedule_exam_success': True,
            'page': 42
        }
        return render(request, 'pages/schedule/view_detail_schedule.html', data)
    else:
        return render(request, 'pages/404.html')


@login_required()
def delete_schedule_exam(request, exam_id, schedule_id):
    exam_schedules = None
    if ExamSchedule.objects.filter(exam_id=exam_id) is not None:
        exam_schedules = ExamSchedule.objects.filter(exam_id=exam_id)
    subject_exam = SubjectExam.objects.get(exam_id=exam_schedules[0].exam.exam_id)
    ExamSchedule.objects.filter(schedule_id=schedule_id).delete()
    data = {
        'exam_schedules': exam_schedules,
        'subject_exam': subject_exam,
        'delete_schedule_exam_success': True,
        'page': 42
    }
    return render(request, 'pages/schedule/view_detail_schedule.html', data)


# FUNCTION -------------------------------------------------------------------------------------------------------------

def __get_subject_list():
    return Subject.objects.all()


def __get_exam_list():
    return Exam.objects.all()


def __get_subject_list_by_user(user_id):
    user = get_object_or_404(User, id=user_id)
    user_subjects = user.usersubject_set.all()

    subjects = []
    for user_subject in user_subjects:
        subjects.append(user_subject.subject)
    return subjects


def __get_exam_list_by_subject(subject_id):
    subject = get_object_or_404(Subject, subject_id=subject_id)
    subject_exams = subject.subjectexam_set.all()

    exams = []
    for subject_exam in subject_exams:
        exams.append(subject_exam.exam)
    return exams


def __get_exam_have_schedule_list_by_subject(subject_id):
    subject = get_object_or_404(Subject, subject_id=subject_id)
    subject_exams = subject.subjectexam_set.all()

    exams = []
    for subject_exam in subject_exams:
        if ExamSchedule.objects.filter(exam_id=subject_exam.exam_id).exists():
            exams.append(subject_exam.exam)
    return exams


def __get_question_list_by_exam(exam_id):
    exam = get_object_or_404(Exam, exam_id=exam_id)
    exam_questions = exam.examquestion_set.all()

    questions = []
    for exam_question in exam_questions:
        question_temp = exam_question.question
        question_temp.answer_list = exam_question.answer_list
        question_temp.answer = exam_question.answer
        questions.append(question_temp)
    return questions


def __convert_answer_list_to_dict(list_answer):
    length = len(list_answer)
    keys = []
    dict_answer = dict()

    for letter in string.ascii_lowercase[:length]:
        keys.append(letter)

    for i in range(len(list_answer)):
        dict_answer[keys[i]] = list_answer[i]

    return dict_answer


def __random_list_value_position(lst):
    n = len(lst)
    # Select two random indices
    index1 = random.randint(0, n - 1)
    index2 = random.randint(0, n - 1)
    # Swap the values at the two random indices
    lst[index1], lst[index2] = lst[index2], lst[index1]
    return lst


def __random_dict_value_position(dictionary):
    values_list = list(dictionary.values())
    return __random_list_value_position(values_list)


def __get_correct_list_answer_new(dict_answer, dict_answer_random, list_correct_answer_key):
    list_correct_answer_value = []
    for k in list_correct_answer_key:
        list_correct_answer_value.append(dict_answer[k])

    list_keys_from_dict = [key for key, value in dict_answer_random.items() if value in list_correct_answer_value]
    return list_keys_from_dict
